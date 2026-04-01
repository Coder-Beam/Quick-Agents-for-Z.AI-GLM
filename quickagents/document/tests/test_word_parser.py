"""
Phase 4 Tests - Word Parser

T019-T022: Test Word document parsing.
"""

import unittest
from pathlib import Path
import tempfile

from quickagents.document.parsers.word_parser import WordParser
from quickagents.document.models import DocumentResult, DocumentTable

FIXTURES = Path(__file__).parent / "fixtures"
SAMPLE_DOCX = FIXTURES / "test_sample.docx"


class TestWordParserInit(unittest.TestCase):
    def test_availability(self):
        parser = WordParser()
        self.assertTrue(parser.is_available())

    def test_parser_name(self):
        self.assertEqual(WordParser.PARSER_NAME, "word")

    def test_supported_formats(self):
        parser = WordParser()
        self.assertTrue(parser.supports_format("docx"))
        self.assertFalse(parser.supports_format("pdf"))


class TestWordTextExtraction(unittest.TestCase):
    """T019: Word text + heading structure extraction"""

    @classmethod
    def setUpClass(cls):
        cls.parser = WordParser()
        cls.result = cls.parser.parse(SAMPLE_DOCX)

    def test_is_document_result(self):
        self.assertIsInstance(self.result, DocumentResult)

    def test_source_format(self):
        self.assertEqual(self.result.source_format, "docx")

    def test_title_from_core_props(self):
        self.assertEqual(self.result.title, "Test Requirements Document")

    def test_raw_text_not_empty(self):
        self.assertTrue(len(self.result.raw_text) > 0)

    def test_raw_text_contains_keywords(self):
        self.assertIn("JWT", self.result.raw_text)
        self.assertIn("Authentication", self.result.raw_text)

    def test_paragraphs_extracted(self):
        self.assertGreater(len(self.result.paragraphs), 0)

    def test_sections_detected(self):
        self.assertGreater(len(self.result.sections), 0)

    def test_heading_levels(self):
        levels = {s.level for s in self.result.sections}
        self.assertIn(1, levels)
        self.assertIn(2, levels)

    def test_section_parent_child(self):
        ids = {s.section_id for s in self.result.sections}
        for sec in self.result.sections:
            if sec.parent_id:
                self.assertIn(sec.parent_id, ids)

    def test_section_content_populated(self):
        has_content = any(s.content for s in self.result.sections)
        self.assertTrue(has_content, "Sections should have content")

    def test_no_errors(self):
        self.assertEqual(len(self.result.errors), 0)

    def test_metadata_has_author(self):
        self.assertIn("author", self.result.metadata)
        self.assertEqual(self.result.metadata["author"], "Test Author")


class TestWordTableExtraction(unittest.TestCase):
    """T020: Word table extraction"""

    @classmethod
    def setUpClass(cls):
        cls.parser = WordParser()
        cls.result = cls.parser.parse(SAMPLE_DOCX)

    def test_tables_found(self):
        self.assertGreater(len(self.result.tables), 0)

    def test_table_headers(self):
        table = self.result.tables[0]
        self.assertEqual(table.headers, ["ID", "Description", "Priority"])

    def test_table_rows(self):
        table = self.result.tables[0]
        self.assertEqual(len(table.rows), 3)

    def test_table_first_row(self):
        table = self.result.tables[0]
        self.assertEqual(table.rows[0][0], "REQ-001")
        self.assertEqual(table.rows[0][2], "High")

    def test_table_to_markdown(self):
        table = self.result.tables[0]
        md = table.to_markdown()
        self.assertIn("| ID |", md)
        self.assertIn("---", md)


class TestWordImageExtraction(unittest.TestCase):
    """T021: Word image extraction"""

    @classmethod
    def setUpClass(cls):
        cls.parser = WordParser()
        cls.result = cls.parser.parse(SAMPLE_DOCX)

    def test_images_is_list(self):
        self.assertIsInstance(self.result.images, list)

    def test_no_images_in_sample(self):
        self.assertEqual(len(self.result.images), 0)

    def test_image_structure(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            self._create_docx_with_image(tmpdir)
            path = Path(tmpdir) / "with_image.docx"
            result = self.parser.parse(path)
            for img in result.images:
                self.assertTrue(hasattr(img, "image_id"))
                self.assertTrue(hasattr(img, "image_type"))
                self.assertGreater(img.page_number, 0)

    @staticmethod
    def _create_docx_with_image(tmpdir):
        from docx import Document
        from docx.shared import Cm
        from PIL import Image as PILImage

        doc = Document()
        doc.add_heading("Doc With Image", level=1)
        doc.add_paragraph("Before image")

        png_path = Path(tmpdir) / "test_img.png"
        img = PILImage.new("RGB", (10, 10), color="red")
        img.save(str(png_path))

        doc.add_picture(str(png_path), width=Cm(2))
        doc.add_paragraph("After image")
        doc.save(str(Path(tmpdir) / "with_image.docx"))


class TestWordEdgeCases(unittest.TestCase):
    def test_file_not_found(self):
        parser = WordParser()
        with self.assertRaises(FileNotFoundError):
            parser.parse(Path("nonexistent.docx"))

    def test_empty_document(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "empty.docx"
            doc = Document()
            doc.save(str(path))

            parser = WordParser()
            result = parser.parse(path)
            self.assertIsInstance(result, DocumentResult)
            self.assertEqual(len(result.sections), 0)
            self.assertEqual(len(result.tables), 0)

    def test_single_heading(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "single.docx"
            doc = Document()
            doc.add_heading("Only Heading", level=1)
            doc.add_paragraph("Some text under heading.")
            doc.save(str(path))

            parser = WordParser()
            result = parser.parse(path)
            self.assertEqual(len(result.sections), 1)
            self.assertEqual(result.sections[0].title, "Only Heading")

    def test_deep_nesting(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "deep.docx"
            doc = Document()
            for level in range(1, 7):
                doc.add_heading(f"Heading Level {level}", level=level)
            doc.save(str(path))

            parser = WordParser()
            result = parser.parse(path)
            self.assertEqual(len(result.sections), 6)
            levels = [s.level for s in result.sections]
            self.assertEqual(levels, [1, 2, 3, 4, 5, 6])

    def test_no_headings(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "no_headings.docx"
            doc = Document()
            doc.add_paragraph("Just plain text.")
            doc.add_paragraph("Another paragraph.")
            doc.save(str(path))

            parser = WordParser()
            result = parser.parse(path)
            self.assertEqual(len(result.sections), 0)
            self.assertGreater(len(result.paragraphs), 0)

    def test_multiple_tables(self):
        from docx import Document
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "tables.docx"
            doc = Document()
            for i in range(3):
                t = doc.add_table(rows=2, cols=2)
                t.rows[0].cells[0].text = f"A{i}"
                t.rows[0].cells[1].text = f"B{i}"
                t.rows[1].cells[0].text = f"C{i}"
                t.rows[1].cells[1].text = f"D{i}"
            doc.save(str(path))

            parser = WordParser()
            result = parser.parse(path)
            self.assertEqual(len(result.tables), 3)

    def test_to_dict_roundtrip(self):
        parser = WordParser()
        result = parser.parse(SAMPLE_DOCX)
        d = result.to_dict()
        self.assertIn("sections", d)
        self.assertIn("tables", d)
        self.assertEqual(d["source_format"], "docx")


if __name__ == "__main__":
    unittest.main()
