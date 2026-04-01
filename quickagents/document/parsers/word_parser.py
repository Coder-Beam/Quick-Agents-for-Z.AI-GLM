"""
Word Parser - Layer 1 parser for Word documents (.docx).

Uses python-docx for text, heading, table, and image extraction.
"""

from pathlib import Path
from typing import List, Dict, Any
import logging

from . import BaseParser
from .utils import find_parent_id, build_structure_tree_stack
from ..models import (
    DocumentResult,
    DocumentSection,
    DocumentTable,
    DocumentImage,
)

logger = logging.getLogger(__name__)

_HEADING_LEVELS = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
    "Heading1": 1, "Heading2": 2, "Heading3": 3,
    "Heading4": 4, "Heading5": 5, "Heading6": 6,
    "Title": 1,
}


class WordParser(BaseParser):
    """Word document parser using python-docx"""

    SUPPORTED_FORMATS = ["docx"]
    REQUIRES_DEPENDENCIES = ["docx"]
    PARSER_NAME = "word"

    def __init__(self):
        super().__init__()
        self._docx = None
        if self._deps_available:
            from docx import Document
            self._docx = Document

    def parse(self, file_path: Path) -> DocumentResult:
        if not self._deps_available:
            raise ImportError(
                "Word parsing requires python-docx. "
                "Install with: pip install quickagents[document]"
            )

        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        logger.info(f"Parsing Word: {file_path}")
        metadata = self._extract_metadata(file_path)

        doc = self._docx(str(file_path))

        title = self._extract_title(doc) or file_path.stem
        sections = self._extract_sections(doc)
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        images = self._extract_images(doc)
        raw_text = self._extract_raw_text(doc)
        structure_tree = build_structure_tree_stack(sections)

        core_props = self._extract_core_properties(doc)

        return DocumentResult(
            source_file=str(file_path),
            source_format="docx",
            title=title,
            sections=sections,
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            formulas=[],
            structure_tree=structure_tree,
            metadata={**metadata, **core_props},
            raw_text=raw_text,
            errors=[],
        )

    def _extract_title(self, doc) -> str:
        """Extract document title from core properties or first heading"""
        if doc.core_properties and doc.core_properties.title:
            return doc.core_properties.title
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name in _HEADING_LEVELS:
                text = para.text.strip()
                if text:
                    return text
        return ""

    def _extract_core_properties(self, doc) -> Dict[str, Any]:
        """Extract core document properties"""
        props = doc.core_properties
        result = {}
        if props.author:
            result["author"] = props.author
        if props.created:
            result["created"] = props.created.isoformat()
        if props.modified:
            result["modified"] = props.modified.isoformat()
        if props.subject:
            result["subject"] = props.subject
        if props.category:
            result["category"] = props.category
        return result

    def _get_heading_level(self, para) -> Optional[int]:
        if not para.style:
            return None
        return _HEADING_LEVELS.get(para.style.name)

    def _extract_sections(self, doc) -> List[DocumentSection]:
        """Extract sections based on heading styles"""
        sections: List[DocumentSection] = []
        counter = 0
        content_buffer: List[str] = []
        current_sec_id = ""

        for para in doc.paragraphs:
            level = self._get_heading_level(para)
            text = para.text.strip()

            if level is not None and text:
                if current_sec_id and content_buffer:
                    for sec in sections:
                        if sec.section_id == current_sec_id:
                            sec.content = "\n".join(content_buffer).strip()
                            break
                    content_buffer = []

                counter += 1
                parent_id = find_parent_id(sections, level)
                sec = DocumentSection(
                    section_id=f"S{counter:03d}",
                    title=text,
                    level=level,
                    page_number=1,
                    parent_id=parent_id,
                )
                sections.append(sec)
                current_sec_id = sec.section_id
            elif text and current_sec_id:
                content_buffer.append(text)

        if current_sec_id and content_buffer:
            for sec in sections:
                if sec.section_id == current_sec_id:
                    sec.content = "\n".join(content_buffer).strip()
                    break

        return sections

    def _extract_paragraphs(self, doc) -> List[str]:
        """Extract non-heading paragraphs"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            level = self._get_heading_level(para)
            if level is None:
                paragraphs.append(text)
        return paragraphs

    def _extract_raw_text(self, doc) -> str:
        """Extract all text from document"""
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts)

    def _extract_tables(self, doc) -> List[DocumentTable]:
        """Extract tables from document"""
        tables: List[DocumentTable] = []
        for idx, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text.strip())
                rows_data.append(row_cells)

            if len(rows_data) < 1:
                continue

            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []

            caption = None
            tbl_element = doc.tables[idx]._tbl
            prev_element = tbl_element.getprevious()
            if prev_element is not None and prev_element.tag.endswith("}p"):
                para_text = "".join(
                    node.text or ""
                    for node in prev_element.iter()
                    if node.tag.endswith("}t")
                ).strip()
                lower = para_text.lower()
                if lower.startswith("table") or lower.startswith("表"):
                    caption = para_text

            tables.append(DocumentTable(
                table_id=f"T{idx + 1:03d}",
                page_number=1,
                caption=caption,
                headers=headers,
                rows=data_rows,
            ))
        return tables

    def _extract_images(self, doc) -> List[DocumentImage]:
        """Extract image metadata from document"""
        images: List[DocumentImage] = []
        img_counter = 0
        seen_rids = set()

        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            if rel.rId in seen_rids:
                continue
            seen_rids.add(rel.rId)
            img_counter += 1

            img_type = "unknown"
            target = rel.target_ref
            if target:
                ext = Path(target).suffix.lower().lstrip(".")
                if ext in ("png", "jpeg", "jpg", "gif", "bmp", "tiff", "svg"):
                    img_type = ext

            images.append(DocumentImage(
                image_id=f"IMG{img_counter:03d}",
                image_type=img_type,
                page_number=1,
                description=f"Image {img_counter} ({target})",
            ))

        return images
